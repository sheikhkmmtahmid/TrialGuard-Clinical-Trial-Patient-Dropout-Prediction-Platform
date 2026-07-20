"""
Extract and log the full text of every documentation file (pdf, docx, doc,
rtf) across data/pds and other data folders. These are the "instructions
and explanations" for the data, read in full, not skimmed.
"""
import warnings
from pathlib import Path
warnings.filterwarnings('ignore')

ROOT = Path(r'D:\Trial Guard\data')
OUT = Path(r'D:\Trial Guard\docs\all_docs_read.txt')


def read_pdf(f):
    import pdfplumber
    text = []
    with pdfplumber.open(f) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            text.append(t)
    return "\n".join(text)


def read_docx(f):
    from docx import Document
    doc = Document(f)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def read_rtf(f):
    try:
        from striprtf.striprtf import rtf_to_text
        with open(f, encoding='utf-8', errors='ignore') as fh:
            return rtf_to_text(fh.read())
    except ImportError:
        return None


if __name__ == '__main__':
    import sys
    subfolder = sys.argv[1] if len(sys.argv) > 1 else ''
    target = ROOT / subfolder if subfolder else ROOT
    files = sorted(target.rglob('*'))
    doc_files = [f for f in files if f.is_file() and f.suffix.lower() in
                 ('.pdf', '.docx', '.doc', '.rtf')]
    print(f"Processing {len(doc_files)} document files under {target}")
    with open(OUT, 'a', encoding='utf-8') as out:
        out.write(f"\n\n========== BATCH: {subfolder or 'ALL'} ({len(doc_files)} files) ==========\n")
        for i, f in enumerate(doc_files):
            out.write(f"\n\n===== {f} =====\n")
            try:
                suffix = f.suffix.lower()
                if suffix == '.pdf':
                    text = read_pdf(f)
                elif suffix == '.docx':
                    text = read_docx(f)
                elif suffix == '.rtf':
                    text = read_rtf(f)
                    if text is None:
                        out.write("SKIPPED: striprtf not installed, needs manual read\n")
                        continue
                elif suffix == '.doc':
                    out.write("FLAGGED: old .doc binary format, cannot auto-extract, needs manual read\n")
                    continue
                else:
                    continue
                out.write(text if text.strip() else "(no extractable text)\n")
            except Exception as e:
                out.write(f"ERROR reading file: {e}\n")
            if (i + 1) % 10 == 0:
                print(f"  {i+1}/{len(doc_files)} done")
    print("done, appended to", OUT)
